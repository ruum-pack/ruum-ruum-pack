from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "app-conductor" / "informe-flujos-app-conductor.docx"
CODIGO_REVISADO = date(2026, 8, 17)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(91, 105, 125)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B8C2CC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(3)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Ruum Ruum - App Conductor | Informe de flujos")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def add_title(doc: Document) -> None:
    p = doc.add_paragraph(style="Title")
    p.add_run("Informe detallado de flujos de la App Conductor")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("Ruum Ruum by Movilia - documento operativo generado desde el checkout actual")
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED

    rows = [
        ("Aplicacion", "@ruum/app-conductor"),
        ("Ubicacion revisada", "C:\\Users\\hmlom\\ruum\\apps\\app-conductor"),
        ("Fecha del informe", date.today().isoformat()),
        ("Corte de revision de codigo", CODIGO_REVISADO.isoformat()),
        ("Version del paquete", "1.0.0"),
        ("Alcance", "Flujos de registro, autenticacion, panel, traslados, evidencia, seguridad, cuenta, ganancias, gastos, soporte, offline, Android y observabilidad."),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2200, 7160])
    set_cell_margins(table)
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
        table.cell(i, 0).paragraphs[0].runs[0].bold = True
    set_repeat_table_header(table.rows[0])


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        hdr.cells[i].text = header
        set_cell_shading(hdr.cells[i], LIGHT_BLUE)
        for run in hdr.cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(11, 37, 69)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_geometry(table, widths)
    set_cell_margins(table)
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    set_cell_margins(table, top=120, bottom=120, start=180, end=180)
    cell = table.cell(0, 0)
    set_repeat_table_header(table.rows[0])
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)


def build() -> None:
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "La App Conductor es la superficie operativa para conductores certificados de Ruum Ruum. "
        "Su responsabilidad principal es guiar la ejecucion de traslados de vehiculos: activar disponibilidad, "
        "aceptar oportunidades, seguir instrucciones por estado, documentar el vehiculo, mantener comunicacion y cerrar la operacion con evidencia auditable."
    )
    add_bullets(doc, [
        "El frontend guia la experiencia, pero las decisiones sensibles dependen de Supabase, RLS, RPCs, Storage privado y servicios compartidos de @ruum/api.",
        "El lenguaje visible favorece 'traslados'; internamente persisten rutas y servicios con nombre historico '/viajes'.",
        "La app ya contempla uso Android con Capacitor, seguimiento de ubicacion, push notifications, cola local de evidencia y un modo de mitigacion para conectividad intermitente.",
        "Los importes al conductor se alimentan desde pasaporte_digital, traslados y payouts_conductor; el informe separa los importes backend de los fallbacks visibles de UI que aun existen.",
    ])

    doc.add_heading("2. Arquitectura y superficies principales", level=1)
    doc.add_paragraph(
        "La aplicacion vive en Next.js App Router con tema oscuro, navegacion global, proveedores de viaje activo, sincronizadores de evidencia y tracking. "
        "El layout monta componentes transversales: LiveRegionProvider, ViajeActivoProvider, SincronizadorEvidenciaOffline, NavegacionConductor, EstadoSincronizacionGlobal, EstadoTrackingGlobal, PushNotificationsBootstrap, VersionGate, OperationalAccessibilityBridge, TextInputUppercaseBridge y OfflineShell."
    )
    add_table(doc, ["Capa", "Responsabilidad", "Archivos representativos"], [
        ["Presentacion Next.js", "Rutas, estados visibles, formularios y acciones guiadas.", "src/app/**"],
        ["Servicios compartidos", "Auth, traslados, conductor, evidencia, ganancias y RPCs.", "@ruum/api/services"],
        ["Contratos compartidos", "Estados, reglas de elegibilidad, validaciones y tipos Supabase.", "@ruum/shared"],
        ["UI compartida", "Tarjetas operativas, botones, badges, montos, stepper y avisos.", "@ruum/ui"],
        ["Movil nativo", "Capacitor Android, camara, geolocalizacion, push y tracking en segundo plano.", "android/**, src/lib/*.ts"],
    ], [1900, 3650, 3810])

    doc.add_heading("3. Mapa de rutas", level=1)
    add_table(doc, ["Ruta", "Flujo", "Proposito"], [
        ["/", "Entrada", "Redirige a /onboarding."],
        ["/onboarding", "Primer uso", "Presenta valor operativo y conduce a login o registro."],
        ["/login", "Autenticacion", "Inicia sesion con Supabase Auth y navega al panel."],
        ["/registro", "Alta conductor", "Crea cuenta, expediente, documentos, consentimientos y solicitud de revision."],
        ["/panel", "Inicio operativo", "Disponibilidad, viaje activo, proximo traslado y oportunidades cercanas."],
        ["/viajes", "Gestion de traslados", "Disponibles, mis viajes, historial, calendario y filtros."],
        ["/viajes/[id]", "Detalle operativo", "Paso actual, accion principal, ruta, contacto, vehiculo, evidencia, soporte y chat."],
        ["/viajes/[id]/evidencia", "Registro del vehiculo", "Checklist fotografico, inspeccion y confirmacion inicial/final."],
        ["/ganancias", "Pagos al conductor", "Resumen e historial desde payouts_conductor."],
        ["/cuenta/*", "Cuenta", "Perfil, documentos, preferencias, datos bancarios, seguridad, soporte y legal."],
        ["/notificaciones", "Centro de notificaciones", "Lista avisos, marca lectura y navega al destino."],
        ["/actualizacion-requerida", "Control de version", "Bloquea versiones incompatibles con backend."],
    ], [1900, 2300, 5160])

    doc.add_heading("4. Flujo de primer uso, autenticacion y recuperacion", level=1)
    add_numbered(doc, [
        "El usuario entra a '/' y la app redirige a '/onboarding'.",
        "Onboarding ofrece avanzar entre pasos, omitir, iniciar sesion o crear cuenta.",
        "Login usa Supabase Auth con email y password; en exito navega a '/panel'.",
        "Recuperacion de password envia enlace, el callback de auth enruta a '/nueva-password' y el cambio de password vuelve al panel.",
        "El middleware refresca la sesion en cada request cuando Supabase esta configurado."
    ])
    add_table(doc, ["Elemento", "Detalle operativo"], [
        ["Sesion", "El middleware usa cookies de @supabase/ssr para mantener sesiones largas."],
        ["Errores", "Los mensajes de Auth se normalizan con traducirErrorAuth/traducirErrorOperativo."],
        ["Rutas sin nav", "Login, registro y onboarding ocultan viaje activo y navegacion principal."],
        ["Salida", "El cierre de sesion limpia estado local sensible, push/tracking cuando aplica y vuelve al acceso."],
    ], [2300, 7060])

    doc.add_heading("5. Flujo de registro de conductor", level=1)
    doc.add_paragraph(
        "El registro es un flujo multi-paso que separa creacion de cuenta, captura de expediente y envio a revision. "
        "La cuenta Auth no equivale a autorizacion operativa: el conductor solo opera cuando el expediente queda aprobado y el estado es activo o modo_prueba_supervisada."
    )
    add_table(doc, ["Paso", "Datos / controles", "Salida"], [
        ["Cuenta", "Telefono, email, password, confirmacion, fortaleza y validaciones compartidas.", "signUp; si no hay sesion, OTP/correo pendiente."],
        ["Identidad y domicilio", "Nombre, apellidos, CURP, CP, estado, ciudad, colonia, calle, numero, referencias y contacto de emergencia.", "Expediente local/remoto actualizado."],
        ["Licencia", "Numero, tipo, vigencia, autorizacion de antecedentes y declaracion de no suspension.", "Datos de licencia normalizados."],
        ["Documentos", "Licencia frente, licencia reverso e identificacion oficial.", "Subida con validacion de tipo permitido."],
        ["Revision y envio", "Consentimientos de terminos, privacidad, antecedentes y declaracion.", "Solicitud enviada a revision."],
    ], [1500, 4300, 3560])
    add_bullets(doc, [
        "El borrador local se limita a campos no sensibles y expira logicamente; no guarda CURP, password, domicilio preciso, licencia, contacto ni archivos.",
        "Con sesion autenticada, el borrador remoto se guarda con debounce de 900 ms y reintento al volver la conexion.",
        "El panel redirige a registro si no existe conductor ni solicitud, o muestra estado de revision si la solicitud ya fue enviada.",
    ])

    doc.add_heading("6. Panel operativo", level=1)
    doc.add_paragraph(
        "El panel es la puerta diaria del conductor. Primero valida sesion, expediente y estado operativo. Si el expediente esta en revision, muestra EstadoRevisionConductor; si hay viaje activo, prioriza PanelActiveTrip; de lo contrario muestra PanelHome."
    )
    add_table(doc, ["Condicion", "Comportamiento"], [
        ["Sin Supabase", "Muestra error operativo y evita simular informacion."],
        ["Sin sesion", "Redirige a /login."],
        ["Sin conductor y sin solicitud", "Redirige a /registro."],
        ["Solicitud incompleta", "Redirige a /registro para continuar."],
        ["Solicitud en revision o conductor no aprobado", "Muestra expediente, documentos y salida segura."],
        ["Conductor operativo", "Carga disponibilidad, viajes disponibles, aceptados, proximo viaje y viaje activo."],
    ], [2800, 6560])
    add_bullets(doc, [
        "La disponibilidad tiene estados disponible, no_disponible y en_viaje.",
        "Cambiar a no_disponible requiere confirmacion; en_viaje no puede manipularse manualmente.",
        "Documento bloqueante abre una alerta y conduce a /cuenta/documentos.",
    ])

    doc.add_heading("7. Flujo de oportunidades y lista de traslados", level=1)
    doc.add_paragraph(
        "La ruta '/viajes' organiza la operacion en vistas: disponibles, mis viajes e historial. Incluye calendario, filtros por fecha/estado, ubicacion aproximada al origen y agrupacion por en-curso/proximos/por-cerrar."
    )
    add_table(doc, ["Vista", "Contenido", "Acciones"], [
        ["Disponibles", "Oportunidades aun no asignadas, ordenables por proximidad si hay ubicacion.", "Ver detalles, aceptar, rechazar con motivo y deshacer durante una ventana breve."],
        ["Mis viajes", "Traslados asignados al conductor agrupados por fase.", "Abrir detalle, continuar siguiente paso o cargar registro si aplica."],
        ["Historial", "Traslados finalizados o en cierre/revision.", "Ver viaje finalizado y estado economico cuando exista."],
    ], [1900, 3900, 3560])
    add_bullets(doc, [
        "Aceptar llama a aceptarViaje y mueve localmente la oportunidad a proximos con estado conductor_asignado.",
        "Rechazar registra evento 'rechazo_oferta_conductor' mediante registrarEvento, no como borrado silencioso.",
        "La elegibilidad usa reglas compartidas esElegibleParaViaje segun tipo de vehiculo y perfil del conductor.",
        "La distancia al origen se etiqueta como aproximada y en linea recta; no promete ETA vial.",
        "La card movil actual conserva fallbacks visuales para hora, distancia y ganancia cuando faltan datos; esos valores deben leerse como placeholder de presentacion, no como contrato operativo.",
    ])

    doc.add_heading("8. Flujo completo de vida de un traslado", level=1)
    add_code_block(doc, [
        "pendiente_de_conductor -> conductor_asignado -> conductor_en_camino_al_origen",
        "-> conductor_en_punto_de_recoleccion -> verificacion_vehiculo_en_proceso",
        "-> evidencia_inicial_en_proceso -> evidencia_inicial_completada -> vehiculo_recibido",
        "-> traslado_en_curso -> llegada_a_destino -> evidencia_final_en_proceso",
        "-> evidencia_final_completada -> entrega_confirmada -> pago_pendiente",
        "-> pago_completado -> servicio_cerrado",
    ])
    add_table(doc, ["Etapa", "Estados principales", "Responsabilidad de la app"], [
        ["Asignacion", "pendiente_de_conductor, conductor_asignado", "Mostrar oportunidad, aceptar y preparar ruta al origen."],
        ["Recoleccion", "conductor_en_camino_al_origen, conductor_en_punto_de_recoleccion", "Guiar llegada, geocerca, contacto de entrega y datos de vehiculo."],
        ["Verificacion inicial", "verificacion_vehiculo_en_proceso, evidencia_inicial_en_proceso", "Capturar fotos, inspeccion y bloquear avance si faltan datos."],
        ["Traslado", "evidencia_inicial_completada, vehiculo_recibido, traslado_en_curso", "Confirmar recepcion, iniciar traslado, seguir ubicacion y llegada al destino."],
        ["Entrega", "llegada_a_destino, evidencia_final_en_proceso, evidencia_final_completada", "Validar contacto de recepcion, evidencia final y comparacion con inicial."],
        ["Cierre", "entrega_confirmada, pago_pendiente, pago_completado, servicio_cerrado", "Cerrar operacion, mostrar revision/pago y permitir disputa si aplica."],
        ["Excepciones", "incidencia_reportada, traslado_fallido, servicio_cancelado, reclamos/disputas", "Detener avance automatico y derivar a Torre de Control/soporte."],
    ], [1700, 2900, 4760])

    doc.add_heading("9. Detalle de traslado y acciones contextuales", level=1)
    doc.add_paragraph(
        "El detalle carga el pasaporte digital desde servidor con sesion real. La presentacion del estado vive en trip-presentation.ts y devuelve etapa, titulo, instruccion, accion primaria, acciones secundarias y proximo paso."
    )
    add_table(doc, ["Accion", "Estado", "Implementacion"], [
        ["Iniciar ruta", "conductor_asignado", "Avanza o abre ruta hacia origen segun presentacion."],
        ["Confirmar llegada origen", "conductor_en_camino_al_origen", "Pantalla DirigeteAOrigen con geocerca y confirmacion fuera de radio."],
        ["Confirmar contacto y vehiculo", "conductor_en_punto_de_recoleccion", "ContactoYVehiculo exige dos confirmaciones antes de evidencias."],
        ["Continuar registro inicial", "evidencia_inicial_en_proceso", "Navega a /viajes/[id]/evidencia."],
        ["Confirmar recepcion", "evidencia_inicial_completada", "Avanza por transicion compartida."],
        ["Iniciar traslado", "vehiculo_recibido", "Inicia etapa de movimiento y tracking."],
        ["Confirmar llegada destino", "traslado_en_curso", "DirigeteADestino usa RPC atomica y geocerca de destino."],
        ["Continuar registro final", "llegada_a_destino / evidencia_final_en_proceso", "Navega a evidencia final y muestra referencia inicial."],
        ["Confirmar entrega / cerrar", "evidencia_final_completada / entrega_confirmada", "Completa cierre operativo y revision."],
        ["Contactar soporte", "incidencia_reportada sin autorizacion", "No avanza estado; requiere decision de Torre de Control."],
    ], [2200, 2600, 4560])
    add_bullets(doc, [
        "El detalle incluye ruta actual, contacto relevante, vehiculo, conteos de evidencia, comparacion inicial/final, reporte de problema, emergencia, disputa y chat.",
        "Las acciones genericas llaman avanzarEstadoTraslado; los pasos con pantalla propia no usan el boton plano para evitar saltos de contexto.",
        "La disputa se habilita solo para cierres o reclamos resueltos dentro de 72 horas desde actualizado_en.",
    ])

    doc.add_heading("10. Evidencia e inspeccion del vehiculo", level=1)
    doc.add_paragraph(
        "La pantalla de evidencia detecta el tipo por estado: inicial para verificacion/evidencia inicial, final para llegada/evidencia final. Combina evidencia remota firmada temporalmente con fotos locales pendientes."
    )
    add_table(doc, ["Requisito fotografico", "Obligatorio", "Uso"], [
        ["Frente", "Si", "Placas y defensa visibles."],
        ["Lado piloto", "Si", "Costado completo del conductor."],
        ["Lado copiloto", "Si", "Costado completo del copiloto."],
        ["Trasera", "Si", "Placa, cajuela y defensa."],
        ["Tablero", "Si", "Odometro/kilometraje legible."],
        ["Dano previo / danos visibles", "No", "Permite no aplica; documenta cambios o hallazgos."],
    ], [3000, 1400, 4960])
    add_table(doc, ["Dato de inspeccion", "Regla"], [
        ["Combustible", "Obligatorio; opciones R, 1/8, 1/4, 3/8, 1/2, 3/4, 1/1."],
        ["Kilometraje", "Obligatorio y numerico no negativo."],
        ["Llaves recibidas", "Obligatorio; opciones 1, 2, 3."],
        ["Holograma de verificacion", "Obligatorio; si/no."],
        ["Talon de verificacion", "Obligatorio."],
        ["Tarjeta de circulacion", "Obligatorio."],
        ["Placa delantera y trasera", "Obligatorias."],
        ["Notas", "Opcional."],
    ], [3000, 6360])
    add_bullets(doc, [
        "Confirmar evidencia guarda inspeccion y llama confirmarEvidenciaCompleta; si faltan campos, devuelve el foco al paso pendiente.",
        "La evidencia se almacena primero en cola local con localId, tipo, angulo, traslado, metadatos, retryCount y errores previos.",
        "La sincronizacion sube blobs a bucket privado 'evidencia', guarda path relativo en evidencia_fotos.url y genera signed URLs solo para visualizar.",
        "El backoff de reintentos escala de 1 minuto a 5, 15 y 60 minutos.",
    ])

    doc.add_heading("11. Incidencias, emergencia, soporte y chat", level=1)
    add_bullets(doc, [
        "Reportar incidencia usa una hoja/modal accesible con foco contenido, tipo de problema, descripcion y evidencia opcional en Storage privado.",
        "EmergencyPanel permite 911, soporte, compartir ubicacion y registrar eventos de accidente o imposibilidad de continuar.",
        "El chat del viaje esta disponible segun reglas compartidas y permite mensajes rapidos; junto al chat existe accion de llamada enmascarada cuando la integracion lo permite.",
        "El soporte de cuenta centraliza telefono, correo/WhatsApp operativo y baja de cuenta; /viajes/[id]/soporte redirige a /cuenta/soporte con contexto de traslado.",
    ])

    doc.add_heading("12. Ganancias y datos bancarios", level=1)
    doc.add_paragraph(
        "El modulo de ganancias consulta datos bancarios, payouts_conductor y traslados del conductor autenticado. La vista pasaporte_digital fue actualizada para calcular ganancia_conductor cuando la ganancia congelada aun es nula, y calcular_pago_conductor ahora devuelve null si no hay precio o certificacion, evitando errores en viajes sin conductor o sin certificacion asignada."
    )
    add_table(doc, ["Flujo", "Detalle"], [
        ["Ganancias", "Carga resumen e historial de pagos; muestra vehiculos trasladados, generado, gastos autorizados, ajustes, retenciones y deposito final cuando existen."],
        ["Fallback economico visible", "La pagina de ganancias aun calcula un monto de respaldo con precio_final/precio_cotizado * 0.85 si no hay ganancia_conductor_congelada; debe tratarse como riesgo de honestidad financiera hasta que se elimine o se respalde por contrato backend."],
        ["Gastos de traslado", "La migracion 20260815000100 habilita que el conductor administre gastos_traslado solo de traslados asignados a su Auth mediante policy RLS."],
        ["Datos bancarios", "Captura titular, banco, CLABE y tarjeta; requiere reautenticacion y llama conductor_guarda_datos_bancarios."],
        ["Auditoria", "El guardado queda en datos_bancarios_conductor con estado en_revision y evento sin exponer numeros completos."],
    ], [2300, 7060])

    doc.add_heading("13. Cuenta, documentos, preferencias y seguridad", level=1)
    add_table(doc, ["Seccion", "Ruta", "Proposito"], [
        ["Perfil", "/cuenta/perfil", "Datos personales, telefono, direccion, emergencia y foto."],
        ["Documentos", "/cuenta/documentos", "Checklist de expediente operativo y reemplazo de documentos."],
        ["Preferencias", "/cuenta/preferencias", "Notificaciones y preferencias de viaje."],
        ["Datos bancarios", "/cuenta/datos-bancarios", "Cuenta de deposito a conductor."],
        ["Seguridad", "/cuenta/seguridad", "Password, sesiones y cambios sensibles."],
        ["Soporte", "/cuenta/soporte", "Canales oficiales, contexto de traslado y baja."],
        ["Legal", "/cuenta/legal", "Terminos y aviso de privacidad."],
    ], [1900, 2300, 5160])

    doc.add_heading("14. Offline, sincronizacion, tracking y Android", level=1)
    add_bullets(doc, [
        "El modo offline cubre evidencia pendiente y cache de viaje activo ya cargado; no equivale a arranque completo sin red.",
        "ViajeActivoProvider consulta el viaje activo cada 45 segundos, al volver visible la app y al reactivar Capacitor; si falla, usa cache local y marca el dato como sin actualizar.",
        "El tracking web reporta ubicacion con minimo de 10 segundos o 50 metros; en Android usa plugin nativo BackgroundTracking y Foreground Service.",
        "PushNotificationsBootstrap registra token, deviceId y aperturas mediante RPCs; las notificaciones dirigen a rutas internas.",
        "Capacitor Android carga la URL remota configurada, con permisos de camara y ubicacion; iOS no esta agregado.",
    ])
    add_table(doc, ["Capacidad", "Implementacion", "Limitacion actual"], [
        ["Camara/evidencia", "@capacitor/camera y fallback de input file", "La evidencia offline depende de sesion/app ya cargada."],
        ["Ubicacion", "Geolocation web/nativa y BackgroundTracking", "Requiere permisos y validacion fisica por dispositivo."],
        ["Push", "@capacitor/push-notifications + RPCs", "Depende de FCM/canal autorizado."],
        ["Arranque sin red", "OfflineShell + cache de viaje activo", "No reconstruye una operacion completa desde cero."],
    ], [2000, 3600, 3760])

    doc.add_heading("15. Controles de seguridad, privacidad y datos reales", level=1)
    add_bullets(doc, [
        "Sin variables NEXT_PUBLIC_SUPABASE_URL/NEXT_PUBLIC_SUPABASE_ANON_KEY la app muestra errores o estados vacios; no simula datos operativos.",
        "El bucket de evidencia es privado; no deben persistirse signed URLs, tokens ni contenido base64 en logs o documentos.",
        "La app evita exponer completos datos sensibles mediante mascaras, tooltips y mensajes de contexto.",
        "Las transiciones sensibles deben permanecer server-owned en RPCs, RLS y servicios compartidos.",
        "Los errores de conectividad se registran con codigos y metadatos operativos, no con payloads sensibles.",
        "La aprobacion operativa no depende de crear cuenta Auth: el panel exige estado_expediente aprobado y conductor activo o modo_prueba_supervisada antes de cargar operacion.",
    ])

    doc.add_heading("16. Accesibilidad y UX operacional", level=1)
    add_bullets(doc, [
        "El layout incluye skip link y LiveRegionProvider para anuncios de estado.",
        "Las pantallas criticas tienen estados de carga, error y vacio; los controles usan labels y ARIA segun el componente.",
        "La navegacion muestra viaje activo y acciones rapidas sin obligar al conductor a recordar rutas.",
        "Las acciones de riesgo, evidencia faltante y decisiones de Torre de Control bloquean el avance con mensajes explicitos.",
    ])

    doc.add_heading("17. Riesgos, pendientes y recomendaciones", level=1)
    add_table(doc, ["Area", "Pendiente / riesgo", "Recomendacion"], [
        ["Offline", "No hay shell local completo para arranque sin red.", "Definir cache persistente de operacion, evidencia/emergencia local y resolucion de conflictos."],
        ["Android", "Matrices fisicas por dispositivo pueden estar pendientes.", "Ejecutar validacion formal de permisos, bateria, cierre forzado, TalkBack y 60 min de ruta."],
        ["Tracking", "Foreground Service existe, pero requiere evidencia de piloto.", "Conservar actas por dispositivo y logs estructurados sin datos sensibles."],
        ["Push", "Depende de configuracion de FCM y canales.", "Validar registro, entrega, apertura y desactivacion por logout."],
        ["Geocerca origen", "Destino tiene RPC atomica; origen podria requerir rigor equivalente.", "Decidir si Operacion necesita RPC atomica de llegada a origen."],
        ["Ganancias", "Hay fallbacks de monto en UI de viajes/ganancias cuando falta ganancia congelada o datos de pasaporte.", "Eliminar inferencias de dinero en frontend o etiquetarlas como estimaciones no pagaderas respaldadas por backend."],
        ["Detalle de traslado", "TripDetailsClient conserva placeholders de distancia, pasajeros, autos y direcciones para ciertos campos vacios.", "Reemplazar por 'Por confirmar' sin numero ficticio o bloquear el dato hasta recibir contrato real."],
        ["Migraciones", "Ambientes compartidos pueden no tener la cadena completa.", "Verificar contratos RPC/RLS antes de pruebas productivas."],
    ], [1800, 3800, 3760])

    doc.add_heading("18. Fuentes revisadas", level=1)
    add_table(doc, ["Archivo / modulo", "Uso en el informe"], [
        ["apps/app-conductor/README.md", "Arquitectura, limites, Android, chat, evidencia privada y pendientes."],
        ["src/app/layout.tsx y middleware.ts", "Proveedores globales, navegacion, sesion y shell transversal."],
        ["src/app/registro/**", "Flujo de alta, pasos, borrador, OTP, documentos y consentimiento."],
        ["src/app/panel/**", "Panel operativo, expediente en revision, disponibilidad y viaje activo."],
        ["src/app/viajes/**", "Vistas, filtros, oportunidades, rechazo, historial y detalle."],
        ["src/lib/trip-presentation.ts", "Mapa de estados a acciones contextuales."],
        ["src/app/viajes/[id]/evidencia/**", "Checklist, inspeccion, comparacion y confirmacion."],
        ["src/lib/cola-offline.ts", "Cola local, sincronizacion, Storage privado y backoff."],
        ["src/app/ViajeActivoContext.tsx y hooks asociados", "Cache, suscripcion, tracking y estado global."],
        ["src/app/ganancias y src/app/cuenta/**", "Pagos, datos bancarios, perfil, documentos y soporte."],
        ["packages/api/src/services/traslados.ts", "Aceptacion, elegibilidad, avance por RPC y llegada atomica a destino."],
        ["packages/api/src/services/evidencia.ts", "Completitud, metodo de pago requerido y auditoria de evidencia inicial/final."],
        ["packages/api/src/services/conductores.ts", "Ganancias, datos bancarios, disponibilidad, historial y preferencias."],
        ["supabase/migrations/20260815000100, 20260816000200, 20260816000300", "Politicas de gastos, ganancia del conductor en pasaporte y correccion de certificacion NULL."],
        ["package.json", "Versiones, scripts y dependencias."],
    ], [3600, 5760])

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.core_properties.title = "Informe detallado de flujos de la App Conductor"
    doc.core_properties.subject = "Ruum Ruum - App Conductor"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Generado desde revision local del repositorio."
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
